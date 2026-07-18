"""Build the checked LibreOffice-produced merged-table PDF corpus fixture.

This generator intentionally uses a producer independent from ReportLab:

    python3 -m pip install \
      -c constraints-pdf-fixtures.txt -e ".[fixtures]"
    SPECTRAIL_SOFFICE=/path/to/soffice \
      python3 tests/fixtures/build_pdf_merged_table_libreoffice_fixture.py

LibreOffice creates the page content. pypdf only rewrites document metadata so
repeated generation with the same LibreOffice build produces stable bytes.
The checked manifest locks the complete fixture toolchain. An intentional
toolchain upgrade additionally requires:

    SPECTRAIL_ACCEPT_FIXTURE_TOOLCHAIN_CHANGE=1
"""

from __future__ import annotations

import hashlib
import json
import os
from pathlib import Path
import shutil
import subprocess
import tempfile

import docx
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from pypdf import PdfReader, PdfWriter, __version__ as PYPDF_VERSION


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "pdf_table_merged_libreoffice.pdf"
MANIFEST = ROOT / "pdf_table_merged_libreoffice.manifest.json"
TOOLCHAIN_CHANGE_ENV = "SPECTRAIL_ACCEPT_FIXTURE_TOOLCHAIN_CHANGE"


def _set_document_layout(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(6)
    section.page_height = Inches(4)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Liberation Sans"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)


def _prepare_table(document: Document):
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        row.height = Inches(0.72)
        for cell in row.cells:
            cell.width = Inches(2.5)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


def _set_cell_text(cell, value: str) -> None:
    cell.text = value
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _build_source_docx(path: Path) -> None:
    document = Document()
    _set_document_layout(document)
    document.core_properties.title = (
        "SpecTrail LibreOffice merged PDF table corpus"
    )
    document.core_properties.author = "SpecTrail"
    document.core_properties.subject = (
        "M5.1 multi-producer merged-table acceptance"
    )

    horizontal = _prepare_table(document)
    horizontal_header = horizontal.cell(0, 0).merge(horizontal.cell(0, 1))
    _set_cell_text(horizontal_header, "LO merged requirement header")
    _set_cell_text(horizontal.cell(1, 0), "REQ-LO-H")
    _set_cell_text(horizontal.cell(1, 1), "Accepted")

    document.add_section(WD_SECTION.NEW_PAGE)
    vertical = _prepare_table(document)
    vertical_owner = vertical.cell(0, 0).merge(vertical.cell(1, 0))
    _set_cell_text(vertical_owner, "LO shared control")
    _set_cell_text(vertical.cell(0, 1), "LO first state")
    _set_cell_text(vertical.cell(1, 1), "LO second state")

    document.save(path)


def _soffice_path() -> str:
    configured = os.environ.get("SPECTRAIL_SOFFICE")
    discovered = configured or shutil.which("soffice") or shutil.which(
        "libreoffice"
    )
    if not discovered:
        raise RuntimeError(
            "LibreOffice is required; set SPECTRAIL_SOFFICE=/path/to/soffice"
        )
    return discovered


def _soffice_identity(executable: str) -> str:
    result = subprocess.run(
        [executable, "--version"],
        check=True,
        capture_output=True,
        text=True,
    )
    identity = result.stdout.strip()
    if not identity:
        raise RuntimeError("LibreOffice did not report a version identity")
    return identity


def _fixture_toolchain(producer_identity: str) -> dict[str, dict[str, str]]:
    return {
        "content_producer": {
            "name": "LibreOffice",
            "identity": producer_identity,
        },
        "source_builder": {
            "name": "python-docx",
            "version": docx.__version__,
        },
        "metadata_normalizer": {
            "name": "pypdf",
            "version": PYPDF_VERSION,
        },
    }


def _check_locked_toolchain(
    toolchain: dict[str, dict[str, str]],
) -> None:
    print(
        "Fixture toolchain: "
        + json.dumps(toolchain, ensure_ascii=False, sort_keys=True)
    )
    if not MANIFEST.exists():
        return
    locked_manifest = json.loads(MANIFEST.read_text(encoding="utf-8"))
    locked = {
        name: locked_manifest.get(name)
        for name in toolchain
    }
    if locked == toolchain:
        return
    if os.environ.get(TOOLCHAIN_CHANGE_ENV) == "1":
        print(
            "Accepting intentional fixture toolchain change because "
            f"{TOOLCHAIN_CHANGE_ENV}=1"
        )
        return
    raise RuntimeError(
        "FIXTURE_TOOLCHAIN_MISMATCH: checked manifest toolchain differs "
        "from the current generator environment; install the locked fixture "
        "dependencies and LibreOffice build, or set "
        f"{TOOLCHAIN_CHANGE_ENV}=1 for an intentional corpus migration. "
        f"locked={locked!r}, current={toolchain!r}"
    )


def _convert_with_libreoffice(
    source: Path,
    output_dir: Path,
    *,
    executable: str,
) -> Path:
    profile = output_dir / "libreoffice-profile"
    result = subprocess.run(
        [
            executable,
            "--headless",
            f"-env:UserInstallation={profile.as_uri()}",
            "--convert-to",
            "pdf",
            "--outdir",
            output_dir.as_posix(),
            source.as_posix(),
        ],
        check=False,
        capture_output=True,
        text=True,
    )
    converted = output_dir / f"{source.stem}.pdf"
    if result.returncode != 0 or not converted.exists():
        raise RuntimeError(
            "LibreOffice PDF conversion failed: "
            f"exit={result.returncode}, stdout={result.stdout!r}, "
            f"stderr={result.stderr!r}"
        )
    return converted


def _normalize_pdf_metadata(
    source: Path,
    target: Path,
    *,
    producer_identity: str,
) -> int:
    reader = PdfReader(source)
    writer = PdfWriter()
    writer.append_pages_from_reader(reader)
    writer.add_metadata(
        {
            "/Title": "SpecTrail LibreOffice merged PDF table corpus",
            "/Author": "SpecTrail",
            "/Subject": "M5.1 multi-producer merged-table acceptance",
            "/Creator": producer_identity,
            "/Producer": (
                f"{producer_identity}; metadata normalized by pypdf"
            ),
            "/CreationDate": "D:20260101000000Z",
            "/ModDate": "D:20260101000000Z",
        }
    )
    with target.open("wb") as stream:
        writer.write(stream)
    return len(reader.pages)


def _write_manifest(
    *,
    toolchain: dict[str, dict[str, str]],
    page_count: int,
) -> None:
    with OUTPUT.open("rb") as stream:
        pdf_sha256 = hashlib.file_digest(stream, "sha256").hexdigest()
    payload = {
        "schema_version": "pdf_fixture_manifest_v2",
        "fixture": OUTPUT.name,
        "pdf_sha256": pdf_sha256,
        "page_count": page_count,
        **toolchain,
        "cases": [
            {
                "page": 1,
                "topology": "horizontal_merge",
                "expected_table": {
                    "row_count": 2,
                    "column_count": 2,
                },
                "expected_cells": [
                    {
                        "row_index": 1,
                        "column_index": 1,
                        "row_span": 1,
                        "column_span": 2,
                        "text": "LO merged requirement header",
                    },
                    {
                        "row_index": 2,
                        "column_index": 1,
                        "row_span": 1,
                        "column_span": 1,
                        "text": "REQ-LO-H",
                    },
                    {
                        "row_index": 2,
                        "column_index": 2,
                        "row_span": 1,
                        "column_span": 1,
                        "text": "Accepted",
                    },
                ],
                "source": {
                    "quote": "LO merged requirement header",
                    "cell_ids": ["cell_00000001_r0001_c0001"],
                    "selected_row_index": 1,
                },
                "projection": {
                    "owner_cell_id": "cell_00000001_r0001_c0001",
                    "span_field": "column_span",
                    "span_value": 2,
                    "occurrence_role": "original",
                },
            },
            {
                "page": 2,
                "topology": "vertical_merge",
                "expected_table": {
                    "row_count": 2,
                    "column_count": 2,
                },
                "expected_cells": [
                    {
                        "row_index": 1,
                        "column_index": 1,
                        "row_span": 2,
                        "column_span": 1,
                        "text": "LO shared control",
                    },
                    {
                        "row_index": 1,
                        "column_index": 2,
                        "row_span": 1,
                        "column_span": 1,
                        "text": "LO first state",
                    },
                    {
                        "row_index": 2,
                        "column_index": 2,
                        "row_span": 1,
                        "column_span": 1,
                        "text": "LO second state",
                    },
                ],
                "source": {
                    "quote": "LO shared control | LO second state",
                    "cell_ids": [
                        "cell_00000002_r0001_c0001",
                        "cell_00000002_r0002_c0002",
                    ],
                    "selected_row_index": 2,
                },
                "projection": {
                    "owner_cell_id": "cell_00000002_r0001_c0001",
                    "span_field": "row_span",
                    "span_value": 2,
                    "occurrence_role": "row_span_projection",
                },
            },
        ],
    }
    MANIFEST.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
    )


def build_fixture() -> None:
    executable = _soffice_path()
    producer_identity = _soffice_identity(executable)
    toolchain = _fixture_toolchain(producer_identity)
    _check_locked_toolchain(toolchain)
    with tempfile.TemporaryDirectory(
        prefix="spectrail-pdf-merged-libreoffice-"
    ) as temporary:
        work = Path(temporary)
        source = work / "pdf_table_merged_libreoffice.docx"
        _build_source_docx(source)
        converted = _convert_with_libreoffice(
            source,
            work,
            executable=executable,
        )
        page_count = _normalize_pdf_metadata(
            converted,
            OUTPUT,
            producer_identity=producer_identity,
        )
    _write_manifest(
        toolchain=toolchain,
        page_count=page_count,
    )


if __name__ == "__main__":
    build_fixture()
