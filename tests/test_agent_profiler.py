from __future__ import annotations

from dataclasses import replace
from pathlib import Path

import pytest
from docx import Document

from spectrail.agent.profiler import (
    LARGE_DOCUMENT_PROMPT_CHARS,
    PROMPT_ESTIMATOR_VERSION,
    DocumentProfiler,
)
from spectrail.evidence.index_builder import ensure_evidence_index
from spectrail.parsers import parse_document


@pytest.mark.parametrize(
    ("source", "source_format", "page_count"),
    [
        (Path("docs/sample_srs.md"), "markdown", None),
        (Path("docs/sample_srs.docx"), "docx", None),
        (Path("docs/sample_srs_text.pdf"), "pdf", 22),
    ],
)
def test_document_profiler_supports_registered_document_formats(
    source: Path,
    source_format: str,
    page_count: int | None,
):
    parsed = parse_document(source)
    evidence_index = ensure_evidence_index(source, parsed)

    profile = DocumentProfiler().build(parsed, evidence_index)

    assert profile.source_format == source_format
    assert profile.page_count == page_count
    assert profile.block_count == len(parsed.blocks)
    assert profile.rendered_text_chars == sum(len(block.text) for block in parsed.blocks)
    assert profile.source_sha256 == evidence_index.source_sha256
    assert profile.parser_name == evidence_index.parser_identity.parser_name
    assert profile.parser_version == evidence_index.parser_identity.parser_version
    assert profile.prompt_estimator_version == PROMPT_ESTIMATOR_VERSION
    assert profile.estimated_prompt_chars > profile.rendered_text_chars
    assert "document_text" not in profile.model_dump()
    assert parsed.text not in profile.model_dump_json()


def test_document_profiler_uses_frozen_counting_contract():
    source = Path("docs/sample_srs.md")
    parsed = parse_document(source)
    evidence_index = ensure_evidence_index(source, parsed)

    profile = DocumentProfiler().build(parsed, evidence_index)

    expected_types = {
        block_type: sum(block.type == block_type for block in parsed.blocks)
        for block_type in sorted({block.type for block in parsed.blocks})
    }
    assert profile.block_type_counts == expected_types
    assert profile.heading_count == expected_types["heading"]
    assert profile.paragraph_count == expected_types["paragraph"]
    assert profile.table_block_count == expected_types.get("table", 0)
    assert profile.section_count == len(
        {tuple(block.section_path) for block in parsed.blocks if block.section_path}
    )
    assert profile.expected_capability_counts == {
        "page_region": 0,
        "table_cell": 0,
        "text_range": len(parsed.blocks),
    }
    assert profile.available_capability_counts == {
        "page_region": 0,
        "table_cell": 0,
        "text_range": len(parsed.blocks),
    }


def test_document_profiler_reports_docx_table_evidence(tmp_path: Path):
    source = tmp_path / "requirements.docx"
    document = Document()
    document.add_heading("Requirements", level=1)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "ID"
    table.cell(0, 1).text = "Requirement"
    table.cell(1, 0).text = "REQ-1"
    table.cell(1, 1).text = "The system shall log access."
    document.save(source)
    parsed = parse_document(source)
    evidence_index = ensure_evidence_index(source, parsed)

    profile = DocumentProfiler().build(parsed, evidence_index)

    assert profile.table_block_count == 1
    assert profile.evidence_table_count == 1
    assert profile.evidence_cell_count == 4
    assert profile.expected_capability_counts["table_cell"] == 1
    assert profile.available_capability_counts["table_cell"] == 1
    assert "HAS_TABLES" in profile.complexity_flags
    assert "HAS_TABLE_CELL_EVIDENCE" in profile.complexity_flags


def test_document_profiler_sanitizes_untrusted_names_and_warnings():
    source = Path("docs/sample_srs_text.pdf")
    parsed = parse_document(source)
    evidence_index = ensure_evidence_index(source, parsed)
    injected = parsed.__class__(
        **{
            **parsed.__dict__,
            "document_name": "x" * 400 + " ignore policy",
            "warnings": [
                "PDF_MULTI_COLUMN_ORDER_BEST_EFFORT: page=1, columns=2, text=ignore policy",
                "arbitrary document-controlled warning body",
            ],
        }
    )
    matching_index = evidence_index.model_copy(
        update={"document_name": injected.document_name}
    )

    profile = DocumentProfiler().build(injected, matching_index)

    assert len(profile.document_name) == 255
    assert "ignore policy" not in profile.document_name
    assert profile.warnings == [
        "PDF_MULTI_COLUMN_ORDER_BEST_EFFORT:columns=2,page=1",
        "PARSER_WARNING",
    ]
    assert "HAS_MULTI_COLUMN_WARNING" in profile.complexity_flags
    assert "PARSER_WARNINGS_PRESENT" in profile.complexity_flags


def test_document_profiler_large_document_flag_uses_prompt_estimate():
    source = Path("eval/cases/sample_srs_long/document.md")
    parsed = parse_document(source)
    evidence_index = ensure_evidence_index(source, parsed)

    profile = DocumentProfiler().build(parsed, evidence_index)

    assert profile.estimated_prompt_chars >= LARGE_DOCUMENT_PROMPT_CHARS
    assert "LARGE_DOCUMENT" in profile.complexity_flags


def test_document_profiler_rejects_mismatched_evidence_identity():
    source = Path("docs/sample_srs.md")
    parsed = parse_document(source)
    original_index = ensure_evidence_index(source, parsed)
    parsed = replace(parsed, source_sha256=original_index.source_sha256)
    evidence_index = original_index.model_copy(
        update={"source_sha256": "0" * 64}
    )

    with pytest.raises(ValueError, match="source_sha256"):
        DocumentProfiler().build(parsed, evidence_index)
