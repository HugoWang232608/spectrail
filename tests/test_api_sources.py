from io import BytesIO
from pathlib import Path
from types import SimpleNamespace

import fitz
import pytest
from fastapi.testclient import TestClient
from docx import Document
from docx.oxml import OxmlElement

from spectrail.core.io import read_json, write_json
from spectrail.evidence import (
    EvidenceIndex,
    finalize_evidence_fingerprint,
    sha256_file,
    sha256_text,
)
from spectrail.evidence.pdf_preview import render_pdf_page
from spectrail.api.app import create_app
from spectrail.parsers.docx_parser import DocxParserV2
from spectrail.parsers.pdf_parser import PdfParserV2
from spectrail.tasks import (
    LegacyEvidenceContinuationRebuildRequiredError,
    LocalTaskStore,
)


def test_api_blocks_returns_completed_task_blocks(api_client: TestClient, completed_api_task: dict):
    task_id = completed_api_task["task_id"]
    fingerprint = _reqir_evidence_fingerprint(api_client, task_id)

    response = api_client.get(_blocks_url(task_id, fingerprint))

    assert response.status_code == 200
    payload = response.json()
    assert payload["task_id"] == task_id
    assert payload["run_generation"] == 1
    assert payload["evidence_fingerprint"] == fingerprint
    assert response.headers["cache-control"] == "private, no-store"
    assert response.headers["x-spectrail-run-generation"] == "1"
    assert len(payload["items"]) > 0
    first = payload["items"][0]
    assert {"block_id", "document_id", "type", "text", "section_path", "order", "metadata"} <= set(first)
    assert "order_index" not in first


def test_api_blocks_before_completion_returns_409(api_client: TestClient):
    created = api_client.post("/api/tasks", json={})
    task_id = created.json()["task_id"]

    response = api_client.get(
        f"/api/tasks/{task_id}/blocks"
        "?expected_run_generation=0"
        f"&expected_evidence_fingerprint={'0' * 64}"
    )

    assert response.status_code == 409
    assert response.json()["detail"]["code"] == "TASK_NOT_COMPLETED"


def test_api_blocks_missing_task_returns_404(api_client: TestClient):
    response = api_client.get(_blocks_url("task_missing", "0" * 64))

    assert response.status_code == 404
    assert response.json()["detail"]["code"] == "TASK_NOT_FOUND"


def test_api_blocks_missing_blocks_file_returns_404(api_client: TestClient, completed_api_task: dict):
    task_id = completed_api_task["task_id"]
    fingerprint = _reqir_evidence_fingerprint(api_client, task_id)
    task_dir = api_client.app.state.task_store.get_task_dir(task_id)
    (task_dir / "parsed" / "blocks.json").unlink()

    response = api_client.get(_blocks_url(task_id, fingerprint))

    assert response.status_code == 404
    assert response.json()["detail"]["code"] == "BLOCKS_NOT_FOUND"


def test_api_blocks_converts_order_index_to_order(api_client: TestClient, completed_api_task: dict):
    task_id = completed_api_task["task_id"]
    fingerprint = _reqir_evidence_fingerprint(api_client, task_id)
    task_dir = api_client.app.state.task_store.get_task_dir(task_id)
    blocks_path = task_dir / "parsed" / "blocks.json"
    blocks = read_json(blocks_path)
    blocks[0]["order_index"] = blocks[0].pop("order")
    write_json(blocks_path, blocks)

    response = api_client.get(_blocks_url(task_id, fingerprint))

    assert response.status_code == 200
    first = response.json()["items"][0]
    assert first["order"] == blocks[0]["order_index"]
    assert "order_index" not in first


def test_api_blocks_rejects_reqir_from_another_evidence_version(
    api_client: TestClient,
    completed_api_task: dict,
):
    task_id = completed_api_task["task_id"]

    response = api_client.get(_blocks_url(task_id, "f" * 64))

    assert response.status_code == 409
    assert response.json()["detail"]["code"] == "EVIDENCE_VERSION_CHANGED"


def test_api_blocks_rejects_content_from_another_evidence_generation(
    api_client: TestClient,
    completed_api_task: dict,
):
    task_id = completed_api_task["task_id"]
    fingerprint = _reqir_evidence_fingerprint(api_client, task_id)
    task_dir = api_client.app.state.task_store.get_task_dir(task_id)
    blocks_path = task_dir / "parsed" / "blocks.json"
    blocks = read_json(blocks_path)
    blocks[0]["text"] = f"{blocks[0]['text']} changed"
    write_json(blocks_path, blocks)

    response = api_client.get(_blocks_url(task_id, fingerprint))

    assert response.status_code == 409
    assert response.json()["detail"]["code"] == "BLOCKS_UNAVAILABLE"


def test_api_blocks_rejects_cross_request_generation_race(
    api_client: TestClient,
    completed_api_task: dict,
):
    task_id = completed_api_task["task_id"]
    fingerprint_a = _reqir_evidence_fingerprint(api_client, task_id)
    task_dir = api_client.app.state.task_store.get_task_dir(task_id)
    blocks_path = task_dir / "parsed" / "blocks.json"
    evidence_path = task_dir / "parsed" / "evidence_index.json"

    blocks_b = read_json(blocks_path)
    blocks_b[0]["text"] = f"{blocks_b[0]['text']} generation B"
    write_json(blocks_path, blocks_b)
    evidence_b = EvidenceIndex.model_validate(read_json(evidence_path))
    first_evidence_block = evidence_b.blocks[0].model_copy(
        update={
            "text_length": len(blocks_b[0]["text"]),
            "text_sha256": sha256_text(blocks_b[0]["text"]),
        }
    )
    evidence_b = evidence_b.model_copy(
        update={
            "blocks": [first_evidence_block, *evidence_b.blocks[1:]],
            "warnings": [*evidence_b.warnings, "generation B"],
        }
    )
    evidence_b = finalize_evidence_fingerprint(evidence_b)
    write_json(evidence_path, evidence_b.model_dump(mode="json"))

    stale_reqir = api_client.get(_blocks_url(task_id, fingerprint_a))

    assert stale_reqir.status_code == 409
    assert stale_reqir.json()["detail"]["code"] == "EVIDENCE_VERSION_CHANGED"

    refreshed_reqir = api_client.get(
        _blocks_url(task_id, evidence_b.evidence_fingerprint)
    )
    assert refreshed_reqir.status_code == 200
    assert refreshed_reqir.json()["items"][0]["text"].endswith("generation B")


def test_task_store_returns_fresh_block_projections_from_immutable_cache(
    api_client: TestClient,
    completed_api_task: dict,
):
    task_id = completed_api_task["task_id"]
    fingerprint = _reqir_evidence_fingerprint(api_client, task_id)
    store = api_client.app.state.task_store

    _, _, first = store.read_blocks(
        task_id,
        expected_evidence_fingerprint=fingerprint,
        expected_run_generation=1,
    )
    original_text = first[0]["text"]
    first[0]["text"] = "caller mutation"

    _, _, second = store.read_blocks(
        task_id,
        expected_evidence_fingerprint=fingerprint,
        expected_run_generation=1,
    )

    assert second[0]["text"] == original_text


def test_api_table_evidence_returns_block_scoped_logical_grid(
    api_client: TestClient,
):
    task_id, parsed = _create_completed_docx_table_task(api_client)
    assert parsed.evidence_index is not None
    table = parsed.evidence_index.tables[0]
    block = parsed.evidence_index.blocks[0]

    response = api_client.get(
        f"/api/tasks/{task_id}/tables/{table.table_id}"
        f"/blocks/{block.block_id}/evidence"
        f"?expected_run_generation=1&expected_evidence_fingerprint={parsed.evidence_index.evidence_fingerprint}"
    )

    assert response.status_code == 200
    assert response.headers["cache-control"] == "private, no-store"
    assert response.headers["x-spectrail-run-generation"] == "1"
    payload = response.json()
    assert payload["schema_version"] == "table_evidence_view_v1"
    assert payload["task_id"] == task_id
    assert payload["evidence_fingerprint"] == (
        parsed.evidence_index.evidence_fingerprint
    )
    assert payload["table_id"] == table.table_id
    assert payload["block_id"] == block.block_id
    assert payload["row_count"] == 2
    assert payload["column_count"] == 3
    assert payload["primary_row_start"] == 1
    assert payload["primary_row_end"] == 2
    assert [row["physical_row_index"] for row in payload["rows"]] == [1, 2]
    header = payload["rows"][0]["cells"][0]
    assert header["cell_id"] == "cell_00000001_r0001_c0001"
    assert header["column_span"] == 2
    assert header["text"] == "Header"
    assert header["occurrences"][0]["occurrence_role"] == "original"


def test_api_table_evidence_reuses_projection_for_pdf_table(
    api_client: TestClient,
):
    task_id, parsed = _create_completed_pdf_table_task(api_client)
    assert parsed.evidence_index is not None
    table = parsed.evidence_index.tables[0]
    block_id = table.block_ids[0]

    response = api_client.get(
        f"/api/tasks/{task_id}/tables/{table.table_id}"
        f"/blocks/{block_id}/evidence"
        f"?expected_run_generation=1&expected_evidence_fingerprint={parsed.evidence_index.evidence_fingerprint}"
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["schema_version"] == "table_evidence_view_v1"
    assert payload["page"] == 1
    assert payload["bbox"] == table.bbox.model_dump(mode="json")
    assert payload["row_count"] == 3
    assert payload["column_count"] == 3
    assert payload["rows"][0]["cells"][0]["is_header"] is True
    assert [cell["text"] for cell in payload["rows"][1]["cells"]] == [
        "REQ-001",
        "Approved within 2 seconds",
        "Safety",
    ]


def test_api_table_evidence_exposes_pdf_continuation_lineage(
    api_client: TestClient,
):
    task_id, parsed = _create_completed_pdf_table_task(
        api_client,
        source=Path("tests/fixtures/pdf_table_continuation.pdf"),
    )
    assert parsed.evidence_index is not None
    root, continued, _ = parsed.evidence_index.tables
    block_id = continued.block_ids[0]

    response = api_client.get(
        f"/api/tasks/{task_id}/tables/{continued.table_id}"
        f"/blocks/{block_id}/evidence"
        f"?expected_run_generation=1&expected_evidence_fingerprint="
        f"{parsed.evidence_index.evidence_fingerprint}"
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["continuation_role"] == "continuation"
    assert payload["continuation_group_id"] == "tblcont_00000001"
    assert payload["continuation_sequence"] == 2
    assert payload["continuation_of_table_id"] == root.table_id
    assert payload["continuation_label"] == "table 1"
    assert payload["continuation_basis"] == (
        "explicit_marker_page_edge_header_match"
    )
    assert payload["continued_header_cell_ids"] == (
        continued.continued_header_cell_ids
    )


def test_task_store_rejects_legacy_pdf_continuation_before_fingerprint_validation(
    api_client: TestClient,
):
    task_id, parsed = _create_completed_pdf_table_task(
        api_client,
        source=Path("tests/fixtures/pdf_table_continuation.pdf"),
    )
    assert parsed.evidence_index is not None
    _, continued, _ = parsed.evidence_index.tables
    task_dir = api_client.app.state.task_store.get_task_dir(task_id)
    _write_legacy_continuation_evidence(task_dir)
    store = api_client.app.state.task_store

    with pytest.raises(
        LegacyEvidenceContinuationRebuildRequiredError,
        match="must be rebuilt with the current parser",
    ):
        store.read_table_evidence(
            task_id,
            table_id=continued.table_id,
            block_id=continued.block_ids[0],
            expected_evidence_fingerprint=(
                parsed.evidence_index.evidence_fingerprint
            ),
            expected_run_generation=1,
        )


@pytest.mark.parametrize("route", ["blocks", "table", "page-preview"])
def test_api_source_routes_require_rebuild_for_legacy_pdf_continuation(
    api_client: TestClient,
    route: str,
):
    task_id, parsed = _create_completed_pdf_table_task(
        api_client,
        source=Path("tests/fixtures/pdf_table_continuation.pdf"),
    )
    assert parsed.evidence_index is not None
    _, continued, _ = parsed.evidence_index.tables
    task_dir = api_client.app.state.task_store.get_task_dir(task_id)
    _write_legacy_continuation_evidence(task_dir)
    fingerprint = parsed.evidence_index.evidence_fingerprint
    if route == "blocks":
        endpoint = _blocks_url(task_id, fingerprint)
    elif route == "table":
        endpoint = (
            f"/api/tasks/{task_id}/tables/{continued.table_id}"
            f"/blocks/{continued.block_ids[0]}/evidence"
            f"?expected_run_generation=1&expected_evidence_fingerprint={fingerprint}"
        )
    else:
        endpoint = _preview_url(
            api_client,
            task_id,
            2,
            expected_evidence_fingerprint=fingerprint,
        )

    response = api_client.get(endpoint)

    assert response.status_code == 409
    assert response.json()["detail"] == {
        "code": "EVIDENCE_LEGACY_CONTINUATION_REBUILD_REQUIRED",
        "message": (
            "legacy PDF table continuation Evidence must be rebuilt "
            "with the current parser"
        ),
    }


def test_api_table_evidence_preserves_repeated_header_projection(
    api_client: TestClient,
):
    task_id, parsed = _create_completed_docx_table_task(
        api_client,
        large=True,
    )
    assert parsed.evidence_index is not None
    table = parsed.evidence_index.tables[0]
    second_block_id = table.block_ids[1]

    response = api_client.get(
        f"/api/tasks/{task_id}/tables/{table.table_id}"
        f"/blocks/{second_block_id}/evidence"
        f"?expected_run_generation=1&expected_evidence_fingerprint={parsed.evidence_index.evidence_fingerprint}"
    )

    assert response.status_code == 200
    rows = response.json()["rows"]
    assert [row["physical_row_index"] for row in rows] == [1, 21, 22]
    assert rows[0]["repeated_header"] is True
    assert rows[0]["cells"][0]["cell_id"] == "cell_00000001_r0001_c0001"
    assert rows[0]["cells"][0]["occurrences"][0]["occurrence_role"] == (
        "repeated_header"
    )
    assert rows[1]["physical_row_index"] == 21


def test_api_table_evidence_rejects_foreign_block(api_client: TestClient):
    task_id, parsed = _create_completed_docx_table_task(api_client)
    assert parsed.evidence_index is not None
    table_id = parsed.evidence_index.tables[0].table_id

    response = api_client.get(
        f"/api/tasks/{task_id}/tables/{table_id}"
        "/blocks/blk_9999/evidence"
        f"?expected_run_generation=1&expected_evidence_fingerprint={parsed.evidence_index.evidence_fingerprint}"
    )

    assert response.status_code == 404
    assert response.json()["detail"]["code"] == "TABLE_EVIDENCE_NOT_FOUND"


def test_api_table_evidence_rejects_stale_fingerprint(api_client: TestClient):
    task_id, parsed = _create_completed_docx_table_task(api_client)
    assert parsed.evidence_index is not None
    task_dir = api_client.app.state.task_store.get_task_dir(task_id)
    evidence_path = task_dir / "parsed" / "evidence_index.json"
    payload = read_json(evidence_path)
    payload["warnings"].append("tampered")
    write_json(evidence_path, payload)
    table = parsed.evidence_index.tables[0]

    response = api_client.get(
        f"/api/tasks/{task_id}/tables/{table.table_id}"
        f"/blocks/{table.block_ids[0]}/evidence"
        f"?expected_run_generation=1&expected_evidence_fingerprint={parsed.evidence_index.evidence_fingerprint}"
    )

    assert response.status_code == 409
    assert response.json()["detail"]["code"] == "TABLE_EVIDENCE_UNAVAILABLE"


def test_api_table_evidence_rejects_reqir_from_another_evidence_version(
    api_client: TestClient,
):
    task_id, parsed = _create_completed_docx_table_task(api_client)
    assert parsed.evidence_index is not None
    table = parsed.evidence_index.tables[0]

    response = api_client.get(
        f"/api/tasks/{task_id}/tables/{table.table_id}"
        f"/blocks/{table.block_ids[0]}/evidence"
        f"?expected_run_generation=1&expected_evidence_fingerprint={'f' * 64}"
    )

    assert response.status_code == 409
    assert response.json()["detail"] == {
        "code": "EVIDENCE_VERSION_CHANGED",
        "message": (
            "ReqIR Evidence fingerprint does not match the current task "
            "EvidenceIndex"
        ),
    }


def test_api_table_evidence_cache_invalidates_when_evidence_changes(
    api_client: TestClient,
):
    task_id, parsed = _create_completed_docx_table_task(api_client)
    assert parsed.evidence_index is not None
    table = parsed.evidence_index.tables[0]
    old_fingerprint = parsed.evidence_index.evidence_fingerprint
    endpoint = (
        f"/api/tasks/{task_id}/tables/{table.table_id}"
        f"/blocks/{table.block_ids[0]}/evidence"
    )

    first = api_client.get(
        f"{endpoint}?expected_run_generation=1&expected_evidence_fingerprint={old_fingerprint}"
    )
    assert first.status_code == 200

    task_dir = api_client.app.state.task_store.get_task_dir(task_id)
    evidence_path = task_dir / "parsed" / "evidence_index.json"
    changed = parsed.evidence_index.model_copy(
        update={"warnings": [*parsed.evidence_index.warnings, "regenerated"]}
    )
    changed = finalize_evidence_fingerprint(changed)
    write_json(evidence_path, changed.model_dump(mode="json"))

    stale_reqir = api_client.get(
        f"{endpoint}?expected_run_generation=1&expected_evidence_fingerprint={old_fingerprint}"
    )
    assert stale_reqir.status_code == 409
    assert stale_reqir.json()["detail"]["code"] == "EVIDENCE_VERSION_CHANGED"

    refreshed_reqir = api_client.get(
        f"{endpoint}?expected_run_generation=1&expected_evidence_fingerprint={changed.evidence_fingerprint}"
    )
    assert refreshed_reqir.status_code == 200
    assert refreshed_reqir.json()["evidence_fingerprint"] == (
        changed.evidence_fingerprint
    )


def test_api_table_evidence_reuses_validated_index_for_unchanged_file(
    api_client: TestClient,
    monkeypatch,
):
    task_id, parsed = _create_completed_docx_table_task(api_client)
    assert parsed.evidence_index is not None
    table = parsed.evidence_index.tables[0]
    endpoint = (
        f"/api/tasks/{task_id}/tables/{table.table_id}"
        f"/blocks/{table.block_ids[0]}/evidence"
        "?expected_run_generation=1&expected_evidence_fingerprint="
        f"{parsed.evidence_index.evidence_fingerprint}"
    )

    assert api_client.get(endpoint).status_code == 200

    def fail_if_revalidated(index) -> None:
        del index
        raise AssertionError("unchanged EvidenceIndex should be served from cache")

    monkeypatch.setattr(
        "spectrail.tasks.store.validate_evidence_fingerprint",
        fail_if_revalidated,
    )
    assert api_client.get(endpoint).status_code == 200


def test_table_evidence_cache_evicts_least_recently_used_task(
    tmp_path: Path,
):
    store = LocalTaskStore(
        tmp_path / "bounded-tasks",
        evidence_cache_max_tasks=2,
    )
    client = TestClient(create_app(task_store=store))
    task_ids: list[str] = []
    for _ in range(3):
        task_id, parsed = _create_completed_docx_table_task(client)
        assert parsed.evidence_index is not None
        table = parsed.evidence_index.tables[0]
        response = client.get(
            f"/api/tasks/{task_id}/tables/{table.table_id}"
            f"/blocks/{table.block_ids[0]}/evidence"
            "?expected_run_generation=1&expected_evidence_fingerprint="
            f"{parsed.evidence_index.evidence_fingerprint}"
        )
        assert response.status_code == 200
        task_ids.append(task_id)

    assert list(store._evidence_cache) == task_ids[-2:]


def test_api_pdf_page_preview_returns_bounded_png(api_client: TestClient):
    task_id = _create_completed_pdf_task(
        api_client,
        width=4000,
        height=1000,
    )

    response = api_client.get(
        _preview_url(api_client, task_id, 1),
        headers={"Origin": "http://localhost:5173"},
    )

    assert response.status_code == 200
    assert response.headers["content-type"] == "image/png"
    assert response.headers["cache-control"] == "private, max-age=300"
    assert response.headers["x-spectrail-preview-width"] == "2000"
    assert response.headers["x-spectrail-preview-height"] == "500"
    assert response.headers["x-spectrail-evidence-fingerprint"] == (
        _current_evidence_fingerprint(api_client, task_id)
    )
    assert "X-Spectrail-Evidence-Fingerprint" in (
        response.headers["access-control-expose-headers"]
    )
    assert response.content.startswith(b"\x89PNG\r\n\x1a\n")


def test_api_pdf_page_preview_rejects_missing_page(api_client: TestClient):
    task_id = _create_completed_pdf_task(api_client)

    response = api_client.get(_preview_url(api_client, task_id, 2))

    assert response.status_code == 404
    assert response.json()["detail"]["code"] == "PAGE_PREVIEW_NOT_FOUND"


def test_api_pdf_page_preview_rejects_reqir_from_previous_generation(
    api_client: TestClient,
):
    task_id = _create_completed_pdf_task(api_client)
    fingerprint_a = _current_evidence_fingerprint(api_client, task_id)
    task_dir = api_client.app.state.task_store.get_task_dir(task_id)

    document_b = fitz.open()
    page_b = document_b.new_page(width=500, height=320)
    page_b.insert_text((40, 60), "Generation B preview evidence")
    (task_dir / "input" / "original.pdf").write_bytes(document_b.tobytes())
    document_b.close()
    parsed_b = PdfParserV2().parse(task_dir / "input" / "original.pdf")
    assert parsed_b.evidence_index is not None
    write_json(
        task_dir / "parsed" / "blocks.json",
        [block.model_dump(mode="json") for block in parsed_b.blocks],
    )
    write_json(
        task_dir / "parsed" / "evidence_index.json",
        parsed_b.evidence_index.model_dump(mode="json"),
    )

    response = api_client.get(
        _preview_url(
            api_client,
            task_id,
            1,
            expected_evidence_fingerprint=fingerprint_a,
        )
    )

    assert response.status_code == 409
    assert response.json()["detail"]["code"] == "EVIDENCE_VERSION_CHANGED"


def test_api_pdf_page_preview_rejects_pdf_not_bound_to_evidence(
    api_client: TestClient,
):
    task_id = _create_completed_pdf_task(api_client)
    fingerprint = _current_evidence_fingerprint(api_client, task_id)
    task_dir = api_client.app.state.task_store.get_task_dir(task_id)
    assert api_client.get(
        _preview_url(
            api_client,
            task_id,
            1,
            expected_evidence_fingerprint=fingerprint,
        )
    ).status_code == 200
    (task_dir / "input" / "original.pdf").write_bytes(b"%PDF changed")

    response = api_client.get(
        _preview_url(
            api_client,
            task_id,
            1,
            expected_evidence_fingerprint=fingerprint,
        )
    )

    assert response.status_code == 409
    assert response.json()["detail"] == {
        "code": "EVIDENCE_VERSION_CHANGED",
        "message": "current PDF does not match the task EvidenceIndex",
    }


def test_api_pdf_page_preview_reuses_hash_for_unchanged_source(
    api_client: TestClient,
    monkeypatch,
):
    task_id = _create_completed_pdf_task(api_client)
    preview_url = _preview_url(api_client, task_id, 1)

    assert api_client.get(preview_url).status_code == 200

    def fail_if_rehashed(path: Path) -> str:
        del path
        raise AssertionError("unchanged PDF should reuse its validated hash")

    monkeypatch.setattr("spectrail.tasks.store.sha256_file", fail_if_rehashed)

    assert api_client.get(preview_url).status_code == 200


def test_api_pdf_page_preview_rejects_source_changed_while_hashing(
    api_client: TestClient,
    monkeypatch,
):
    task_id = _create_completed_pdf_task(api_client)
    task_dir = api_client.app.state.task_store.get_task_dir(task_id)
    document_path = task_dir / "input" / "original.pdf"
    preview_url = _preview_url(api_client, task_id, 1)
    assert api_client.get(preview_url).status_code == 200
    cache_entry = api_client.app.state.task_store._evidence_cache[task_id]
    cached_sha256 = cache_entry.source_sha256
    assert cached_sha256 == sha256_file(document_path)

    document_path.write_bytes(b"%PDF generation B")

    def hash_then_change(path: Path) -> str:
        generation_b_sha256 = sha256_file(path)
        Path(path).write_bytes(b"%PDF generation C changed while hashing")
        return generation_b_sha256

    monkeypatch.setattr(
        "spectrail.tasks.store.sha256_file",
        hash_then_change,
    )

    response = api_client.get(preview_url)

    assert response.status_code == 409
    assert response.json()["detail"] == {
        "code": "PAGE_PREVIEW_UNAVAILABLE",
        "message": f"PDF preview source changed while hashing: {task_id}",
    }
    assert cached_sha256 is not None
    assert cache_entry.source_file_signature is None
    assert cache_entry.source_sha256 is None


def test_api_pdf_page_preview_discards_source_changed_while_rendering(
    api_client: TestClient,
    monkeypatch,
):
    task_id = _create_completed_pdf_task(api_client)

    def render_then_change(
        document_path: Path,
        page_number: int,
    ) -> tuple[bytes, int, int]:
        rendered = render_pdf_page(document_path, page_number)
        document_path.write_bytes(b"%PDF changed while rendering")
        return rendered

    monkeypatch.setattr(
        "spectrail.tasks.store.render_pdf_page",
        render_then_change,
    )

    response = api_client.get(_preview_url(api_client, task_id, 1))

    assert response.status_code == 409
    assert response.json()["detail"] == {
        "code": "PAGE_PREVIEW_UNAVAILABLE",
        "message": f"PDF preview source changed while rendering: {task_id}",
    }
    cache_entry = api_client.app.state.task_store._evidence_cache[task_id]
    assert cache_entry.source_file_signature is None
    assert cache_entry.source_sha256 is None


def test_api_pdf_page_preview_preserves_not_found_when_close_fails(
    api_client: TestClient,
    monkeypatch,
):
    task_id = _create_completed_pdf_task(api_client)

    class MissingPageDocument:
        page_count = 0

        def close(self):
            raise RuntimeError("close failed")

    monkeypatch.setattr(fitz, "open", lambda path: MissingPageDocument())

    response = api_client.get(_preview_url(api_client, task_id, 1))

    assert response.status_code == 404
    assert response.json()["detail"] == {
        "code": "PAGE_PREVIEW_NOT_FOUND",
        "message": "PDF page does not exist: 1",
    }


@pytest.mark.parametrize(
    ("rotation", "expected_size"),
    [
        (0, ("800", "600")),
        (90, ("600", "800")),
        (180, ("800", "600")),
        (270, ("600", "800")),
    ],
)
def test_api_pdf_page_preview_uses_rotated_page_dimensions(
    api_client: TestClient,
    rotation: int,
    expected_size: tuple[str, str],
):
    task_id = _create_completed_pdf_task(
        api_client,
        width=400,
        height=300,
        rotation=rotation,
    )

    response = api_client.get(_preview_url(api_client, task_id, 1))

    assert response.status_code == 200
    assert response.headers["x-spectrail-preview-width"] == expected_size[0]
    assert response.headers["x-spectrail-preview-height"] == expected_size[1]
    assert response.headers["x-spectrail-run-generation"] == "1"


@pytest.mark.parametrize(
    "route",
    ["reqir", "blocks", "table", "page-preview"],
)
def test_api_evidence_reads_reject_stale_task_run_generation(
    api_client: TestClient,
    route: str,
):
    task_id, parsed = _create_completed_pdf_table_task(api_client)
    assert parsed.evidence_index is not None
    table = parsed.evidence_index.tables[0]
    fingerprint = parsed.evidence_index.evidence_fingerprint
    if route == "reqir":
        url = (
            f"/api/tasks/{task_id}/reqir"
            "?expected_run_generation=0"
        )
    elif route == "blocks":
        url = (
            f"/api/tasks/{task_id}/blocks"
            "?expected_run_generation=0"
            f"&expected_evidence_fingerprint={fingerprint}"
        )
    elif route == "table":
        url = (
            f"/api/tasks/{task_id}/tables/{table.table_id}"
            f"/blocks/{table.block_ids[0]}/evidence"
            "?expected_run_generation=0"
            f"&expected_evidence_fingerprint={fingerprint}"
        )
    else:
        url = (
            f"/api/tasks/{task_id}/pages/1/preview.png"
            "?expected_run_generation=0"
            f"&expected_evidence_fingerprint={fingerprint}"
        )

    response = api_client.get(url)

    assert response.status_code == 409
    assert response.json()["detail"]["code"] == "RUN_GENERATION_CHANGED"


def test_api_pdf_page_preview_maps_render_failure_to_structured_error(
    api_client: TestClient,
    monkeypatch,
):
    task_id = _create_completed_pdf_task(api_client)

    class BrokenPage:
        rect = SimpleNamespace(width=400, height=300)

        def get_pixmap(self, **kwargs):
            del kwargs
            raise RuntimeError("broken page resource")

    class BrokenDocument:
        page_count = 1

        def __getitem__(self, index):
            assert index == 0
            return BrokenPage()

        def close(self):
            return None

    monkeypatch.setattr(fitz, "open", lambda path: BrokenDocument())

    response = api_client.get(_preview_url(api_client, task_id, 1))

    assert response.status_code == 409
    assert response.json()["detail"] == {
        "code": "PAGE_PREVIEW_UNAVAILABLE",
        "message": "failed to render PDF page preview: 1",
    }


def test_api_pdf_page_preview_preserves_render_error_when_close_fails(
    api_client: TestClient,
    monkeypatch,
    caplog,
):
    task_id = _create_completed_pdf_task(api_client)

    class BrokenPage:
        rect = SimpleNamespace(width=400, height=300)

        def get_pixmap(self, **kwargs):
            del kwargs
            raise RuntimeError("broken page resource")

    class BrokenDocument:
        page_count = 1

        def __getitem__(self, index):
            assert index == 0
            return BrokenPage()

        def close(self):
            raise RuntimeError("close failed")

    monkeypatch.setattr(fitz, "open", lambda path: BrokenDocument())
    caplog.set_level("WARNING", logger="spectrail.evidence.pdf_preview")

    response = api_client.get(_preview_url(api_client, task_id, 1))

    assert response.status_code == 409
    assert response.json()["detail"] == {
        "code": "PAGE_PREVIEW_UNAVAILABLE",
        "message": "failed to render PDF page preview: 1",
    }
    assert "failed to close PDF page preview source 1 after" in caplog.text


def test_api_page_preview_rejects_non_pdf_task(
    api_client: TestClient,
    completed_api_task: dict,
):
    task_id = completed_api_task["task_id"]

    response = api_client.get(
        f"/api/tasks/{task_id}/pages/1/preview.png"
        f"?expected_run_generation=1&expected_evidence_fingerprint={_reqir_evidence_fingerprint(api_client, task_id)}"
    )

    assert response.status_code == 409
    assert response.json()["detail"]["code"] == "PAGE_PREVIEW_UNAVAILABLE"


def test_api_page_preview_before_completion_returns_409(api_client: TestClient):
    created = api_client.post("/api/tasks", json={})
    task_id = created.json()["task_id"]

    response = api_client.get(
        f"/api/tasks/{task_id}/pages/1/preview.png"
        f"?expected_run_generation=0&expected_evidence_fingerprint={'0' * 64}"
    )

    assert response.status_code == 409
    assert response.json()["detail"]["code"] == "TASK_NOT_COMPLETED"


def _create_completed_pdf_task(
    api_client: TestClient,
    *,
    width: float = 400,
    height: float = 300,
    rotation: int = 0,
) -> str:
    created = api_client.post("/api/tasks", json={})
    assert created.status_code == 200
    task_id = created.json()["task_id"]

    document = fitz.open()
    page = document.new_page(width=width, height=height)
    page.insert_text((30, 40), "Preview evidence")
    page.set_rotation(rotation)
    content = document.tobytes()
    document.close()
    uploaded = api_client.post(
        f"/api/tasks/{task_id}/documents",
        files={"file": ("preview.pdf", content, "application/pdf")},
    )
    assert uploaded.status_code == 200
    task_dir = api_client.app.state.task_store.get_task_dir(task_id)
    parsed = PdfParserV2().parse(task_dir / "input" / "original.pdf")
    assert parsed.evidence_index is not None
    write_json(
        task_dir / "parsed" / "blocks.json",
        [block.model_dump(mode="json") for block in parsed.blocks],
    )
    write_json(
        task_dir / "parsed" / "evidence_index.json",
        parsed.evidence_index.model_dump(mode="json"),
    )
    api_client.app.state.task_store.begin_run(task_id)
    api_client.app.state.task_store.update_task(task_id, status="completed")
    return task_id


def _create_completed_docx_table_task(
    api_client: TestClient,
    *,
    large: bool = False,
):
    document = Document()
    if large:
        table = document.add_table(rows=22, cols=2)
        table.cell(0, 0).text = "Key"
        table.cell(0, 1).text = "Requirement"
        _mark_repeating_header(table.rows[0]._tr)
        for row_index in range(1, 22):
            table.cell(row_index, 0).text = f"R{row_index}"
            table.cell(row_index, 1).text = f"Requirement {row_index}"
    else:
        table = document.add_table(rows=2, cols=3)
        table.cell(0, 0).merge(table.cell(0, 1)).text = "Header"
        table.cell(0, 2).text = "Status"
        table.cell(1, 0).text = "REQ-1"
        table.cell(1, 1).text = "The system shall respond."
        table.cell(1, 2).text = "Approved"
    content = BytesIO()
    document.save(content)

    created = api_client.post("/api/tasks", json={})
    task_id = created.json()["task_id"]
    uploaded = api_client.post(
        f"/api/tasks/{task_id}/documents",
        files={
            "file": (
                "table.docx",
                content.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert uploaded.status_code == 200

    task_dir = api_client.app.state.task_store.get_task_dir(task_id)
    parsed = DocxParserV2().parse(task_dir / "input" / "original.docx")
    assert parsed.evidence_index is not None
    write_json(
        task_dir / "parsed" / "blocks.json",
        [block.model_dump(mode="json") for block in parsed.blocks],
    )
    write_json(
        task_dir / "parsed" / "evidence_index.json",
        parsed.evidence_index.model_dump(mode="json"),
    )
    api_client.app.state.task_store.begin_run(task_id)
    api_client.app.state.task_store.update_task(task_id, status="completed")
    return task_id, parsed


def _create_completed_pdf_table_task(
    api_client: TestClient,
    *,
    source: Path = Path("tests/fixtures/pdf_table_requirements.pdf"),
):
    created = api_client.post("/api/tasks", json={})
    task_id = created.json()["task_id"]
    uploaded = api_client.post(
        f"/api/tasks/{task_id}/documents",
        files={
            "file": (
                source.name,
                source.read_bytes(),
                "application/pdf",
            )
        },
    )
    assert uploaded.status_code == 200

    task_dir = api_client.app.state.task_store.get_task_dir(task_id)
    parsed = PdfParserV2().parse(task_dir / "input" / "original.pdf")
    assert parsed.evidence_index is not None
    write_json(
        task_dir / "parsed" / "blocks.json",
        [block.model_dump(mode="json") for block in parsed.blocks],
    )
    write_json(
        task_dir / "parsed" / "evidence_index.json",
        parsed.evidence_index.model_dump(mode="json"),
    )
    api_client.app.state.task_store.begin_run(task_id)
    api_client.app.state.task_store.update_task(task_id, status="completed")
    return task_id, parsed


def _write_legacy_continuation_evidence(task_dir: Path) -> None:
    evidence_path = task_dir / "parsed" / "evidence_index.json"
    legacy_payload = read_json(evidence_path)
    for table_payload in legacy_payload["tables"]:
        table_payload.pop("continuation_basis")
        table_payload.pop("continuation_label")
    write_json(evidence_path, legacy_payload)


def _mark_repeating_header(row) -> None:
    table_header = OxmlElement("w:tblHeader")
    table_header.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val",
        "true",
    )
    row.get_or_add_trPr().append(table_header)


def _reqir_evidence_fingerprint(
    api_client: TestClient,
    task_id: str,
) -> str:
    response = api_client.get(
        f"/api/tasks/{task_id}/reqir?expected_run_generation=1"
    )
    assert response.status_code == 200
    return response.json()["metadata"]["evidence_fingerprint"]


def _blocks_url(task_id: str, evidence_fingerprint: str) -> str:
    return (
        f"/api/tasks/{task_id}/blocks"
        f"?expected_run_generation=1&expected_evidence_fingerprint={evidence_fingerprint}"
    )


def _current_evidence_fingerprint(
    api_client: TestClient,
    task_id: str,
) -> str:
    task_dir = api_client.app.state.task_store.get_task_dir(task_id)
    return read_json(
        task_dir / "parsed" / "evidence_index.json"
    )["evidence_fingerprint"]


def _preview_url(
    api_client: TestClient,
    task_id: str,
    page: int,
    *,
    expected_evidence_fingerprint: str | None = None,
) -> str:
    fingerprint = (
        expected_evidence_fingerprint
        or _current_evidence_fingerprint(api_client, task_id)
    )
    return (
        f"/api/tasks/{task_id}/pages/{page}/preview.png"
        f"?expected_run_generation=1&expected_evidence_fingerprint={fingerprint}"
    )
