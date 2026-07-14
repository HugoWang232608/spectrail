from pathlib import Path
import socket

import fitz
from docx import Document
from fastapi.testclient import TestClient

import spectrail.api.routes.review as review_routes
from spectrail.core.io import write_json
from spectrail.parsers.markdown_parser import MarkdownParser
from spectrail.task_transactions import TaskTransactionError, task_lock


def test_api_task_flow(api_client: TestClient):
    assert api_client.get("/api/health").json() == {"status": "ok"}

    created = api_client.post("/api/tasks", json={"goal": "extract_requirements", "model_mode": "mock"})
    assert created.status_code == 200
    task = created.json()
    task_id = task["task_id"]
    assert task["status"] == "created"

    sample = Path("docs/sample_srs.md")
    uploaded = api_client.post(
        f"/api/tasks/{task_id}/documents",
        files={"file": (sample.name, sample.read_bytes(), "text/markdown")},
    )
    assert uploaded.status_code == 200
    assert uploaded.json()["status"] == "uploaded"
    assert uploaded.json()["filename"] == sample.name

    run = api_client.post(f"/api/tasks/{task_id}/run")
    assert run.status_code == 200
    payload = run.json()
    assert payload["status"] == "completed"
    assert payload["manifest"]["counts"]["validated_requirements"] >= 14

    status = api_client.get(f"/api/tasks/{task_id}")
    assert status.status_code == 200
    assert status.json()["manifest"]["status"] == "completed"

    reqir = api_client.get(f"/api/tasks/{task_id}/reqir")
    assert reqir.status_code == 200
    assert len(reqir.json()["items"]) >= 14


def test_api_forced_chunking_exposes_chunk_and_quarantine_artifacts(api_client: TestClient):
    created = api_client.post(
        "/api/tasks",
        json={
            "model_mode": "mock",
            "chunking_mode": "force",
            "max_rendered_prompt_chars": 1600,
            "overlap_blocks": 1,
            "validation_policy": "quarantine",
            "evidence_policy": "structured_required",
        },
    )
    assert created.status_code == 200
    task_id = created.json()["task_id"]

    sample = Path("docs/sample_srs.md")
    uploaded = api_client.post(
        f"/api/tasks/{task_id}/documents",
        files={"file": (sample.name, sample.read_bytes(), "text/markdown")},
    )
    assert uploaded.status_code == 200

    run = api_client.post(f"/api/tasks/{task_id}/run")
    assert run.status_code == 200
    assert run.json()["manifest"]["counts"]["chunks"] >= 3

    chunks = api_client.get(f"/api/tasks/{task_id}/chunks")
    quarantined = api_client.get(f"/api/tasks/{task_id}/quarantined")
    assert chunks.status_code == 200
    assert len(chunks.json()) >= 3
    assert quarantined.status_code == 200
    assert quarantined.json()["items"] == []


def test_api_rejects_invalid_chunking_configuration(api_client: TestClient):
    response = api_client.post(
        "/api/tasks",
        json={"max_rendered_prompt_chars": 999, "overlap_blocks": 6},
    )
    assert response.status_code == 422

    response = api_client.post(
        "/api/tasks",
        json={"evidence_policy": "unsupported"},
    )
    assert response.status_code == 422


def test_completed_with_warnings_remains_readable_reviewable_and_exportable(
    api_client: TestClient, completed_api_task: dict
):
    task_id = completed_api_task["task_id"]
    api_client.app.state.task_store.update_task(task_id, status="completed_with_warnings")
    reqir = api_client.get(f"/api/tasks/{task_id}/reqir")
    assert reqir.status_code == 200
    requirement_id = reqir.json()["items"][0]["id"]
    reviewed = api_client.post(
        f"/api/tasks/{task_id}/review",
        json={"requirement_id": requirement_id, "action": "approve"},
    )
    exported = api_client.get(f"/api/tasks/{task_id}/exports/requirements.xlsx")
    assert reviewed.status_code == 200
    assert exported.status_code == 200


def test_api_task_requires_uploaded_document(api_client: TestClient):
    created = api_client.post("/api/tasks", json={})
    task_id = created.json()["task_id"]

    run = api_client.post(f"/api/tasks/{task_id}/run")
    assert run.status_code == 409
    assert run.json()["detail"]["code"] == "DOCUMENT_NOT_UPLOADED"


def test_api_reads_and_writes_reject_incomplete_migration(
    api_client: TestClient,
    completed_api_task: dict,
):
    task_id = completed_api_task["task_id"]
    store = api_client.app.state.task_store
    task_dir = store.get_task_dir(task_id)
    (task_dir / ".migration_tmp").mkdir()

    for response in (
        api_client.get(f"/api/tasks/{task_id}"),
        api_client.get(f"/api/tasks/{task_id}/reqir"),
        api_client.get(f"/api/tasks/{task_id}/exports/reqir.json"),
        api_client.post(
            f"/api/tasks/{task_id}/review",
            json={"requirement_id": "REQ-0001", "action": "approve"},
        ),
        api_client.post(f"/api/tasks/{task_id}/run"),
    ):
        assert response.status_code == 409
        assert response.json()["detail"]["code"] == "TASK_MIGRATION_INCOMPLETE"


def test_api_distinguishes_active_task_lock_from_incomplete_migration(
    api_client: TestClient,
    completed_api_task: dict,
):
    task_id = completed_api_task["task_id"]
    task_dir = api_client.app.state.task_store.get_task_dir(task_id)

    with task_lock(task_dir, operation="background_pipeline"):
        response = api_client.get(f"/api/tasks/{task_id}")

    assert response.status_code == 409
    assert response.json()["detail"]["code"] == "TASK_TRANSACTION_LOCKED"
    assert response.json()["detail"]["retryable"] is True


def test_api_operation_reclaims_dead_same_host_process_lock(
    api_client: TestClient,
    completed_api_task: dict,
):
    task_id = completed_api_task["task_id"]
    task_dir = api_client.app.state.task_store.get_task_dir(task_id)
    lock_dir = task_dir / ".task.lock"
    lock_dir.mkdir()
    write_json(
        lock_dir / "owner.json",
        {
            "schema_version": "task_lock_v1",
            "token": "crashed-process",
            "operation": "api_pipeline_run",
            "pid": 99999999,
            "host": socket.gethostname(),
            "started_at": "2026-07-14T00:00:00Z",
        },
    )

    response = api_client.get(f"/api/tasks/{task_id}")

    assert response.status_code == 200
    assert not lock_dir.exists()


def test_review_race_preserves_transaction_error_code(
    api_client: TestClient,
    completed_api_task: dict,
    monkeypatch,
):
    task_id = completed_api_task["task_id"]

    def migration_won_race(**kwargs):
        del kwargs
        raise TaskTransactionError(
            "TASK_MIGRATION_INCOMPLETE",
            "migration started after readability check",
        )

    monkeypatch.setattr(
        review_routes,
        "apply_review_to_package",
        migration_won_race,
    )
    response = api_client.post(
        f"/api/tasks/{task_id}/review",
        json={"requirement_id": "REQ-0001", "action": "approve"},
    )

    assert response.status_code == 409
    assert response.json()["detail"]["code"] == "TASK_MIGRATION_INCOMPLETE"


def test_api_upload_accepts_docx_and_pdf(api_client: TestClient):
    created = api_client.post("/api/tasks", json={})
    task_id = created.json()["task_id"]

    uploaded_docx = api_client.post(
        f"/api/tasks/{task_id}/documents",
        files={
            "file": (
                "sample.docx",
                b"docx bytes",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert uploaded_docx.status_code == 200
    assert uploaded_docx.json()["filename"] == "sample.docx"

    uploaded_pdf = api_client.post(
        f"/api/tasks/{task_id}/documents",
        files={"file": ("sample.pdf", b"%PDF", "application/pdf")},
    )
    assert uploaded_pdf.status_code == 200
    assert uploaded_pdf.json()["filename"] == "sample.pdf"


def test_api_run_docx_task_completed(api_client: TestClient, tmp_path: Path):
    document_path = tmp_path / "sample_srs.docx"
    _write_docx_from_sample_blocks(document_path)

    task_id = _create_and_upload(
        api_client,
        document_path,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    run = api_client.post(f"/api/tasks/{task_id}/run")
    assert run.status_code == 200
    assert run.json()["status"] == "completed"
    assert run.json()["manifest"]["counts"]["validated_requirements"] >= 14

    reqir = api_client.get(f"/api/tasks/{task_id}/reqir")
    assert reqir.status_code == 200
    assert len(reqir.json()["items"]) >= 14

    xlsx = api_client.get(f"/api/tasks/{task_id}/exports/requirements.xlsx")
    assert xlsx.status_code == 200
    assert xlsx.content


def test_api_run_text_pdf_task_completed_with_source_pages(api_client: TestClient, tmp_path: Path):
    document_path = tmp_path / "sample_srs_text.pdf"
    _write_pdf_from_sample_blocks(document_path)

    task_id = _create_and_upload(api_client, document_path, "application/pdf")

    run = api_client.post(f"/api/tasks/{task_id}/run")
    assert run.status_code == 200
    assert run.json()["status"] == "completed"
    assert run.json()["manifest"]["counts"]["validated_requirements"] >= 14

    reqir = api_client.get(f"/api/tasks/{task_id}/reqir")
    blocks = api_client.get(f"/api/tasks/{task_id}/blocks")
    xlsx = api_client.get(f"/api/tasks/{task_id}/exports/requirements.xlsx")
    assert reqir.status_code == 200
    assert blocks.status_code == 200
    assert xlsx.status_code == 200

    block_index = {block["block_id"]: block for block in blocks.json()["items"]}
    assert any(item["sources"][0]["page"] is not None for item in reqir.json()["items"])
    for item in reqir.json()["items"]:
        source = item["sources"][0]
        assert source["page"] == block_index[source["block_id"]]["page"]


def test_api_upload_rejects_unsupported_extension(api_client: TestClient):
    created = api_client.post("/api/tasks", json={})
    task_id = created.json()["task_id"]

    uploaded = api_client.post(
        f"/api/tasks/{task_id}/documents",
        files={"file": ("sample.rtf", b"{\\rtf1}", "application/rtf")},
    )
    assert uploaded.status_code == 400
    assert uploaded.json()["detail"]["code"] == "INVALID_DOCUMENT"


def test_api_run_marks_task_failed_when_model_mode_is_rejected(api_client: TestClient):
    created = api_client.post("/api/tasks", json={})
    task_id = created.json()["task_id"]

    sample = Path("docs/sample_srs.md")
    uploaded = api_client.post(
        f"/api/tasks/{task_id}/documents",
        files={"file": (sample.name, sample.read_bytes(), "text/markdown")},
    )
    assert uploaded.status_code == 200

    api_client.app.state.task_store.update_task(task_id, model_mode="unknown")

    run = api_client.post(f"/api/tasks/{task_id}/run")
    assert run.status_code == 400
    assert run.json()["detail"]["code"] == "INVALID_MODEL_MODE"

    status = api_client.get(f"/api/tasks/{task_id}")
    assert status.status_code == 200
    assert status.json()["status"] == "failed"


def _create_and_upload(api_client: TestClient, document_path: Path, media_type: str) -> str:
    created = api_client.post("/api/tasks", json={})
    task_id = created.json()["task_id"]
    uploaded = api_client.post(
        f"/api/tasks/{task_id}/documents",
        files={"file": (document_path.name, document_path.read_bytes(), media_type)},
    )
    assert uploaded.status_code == 200
    return task_id


def _write_docx_from_sample_blocks(path: Path) -> None:
    document = Document()
    for block in MarkdownParser().parse_file("docs/sample_srs.md"):
        if block.type == "heading":
            document.add_heading(block.text, level=int(block.metadata.get("level", 1)))
        elif block.type == "list":
            document.add_paragraph(block.text, style="List Bullet")
        else:
            document.add_paragraph(block.text)
    document.save(path)


def _write_pdf_from_sample_blocks(path: Path) -> None:
    document = fitz.open()
    for block in MarkdownParser().parse_file("docs/sample_srs.md"):
        page = document.new_page(width=2000, height=400)
        page.insert_text((72, 72), block.text, fontsize=8, fontname="china-s")
    document.save(path)
    document.close()
