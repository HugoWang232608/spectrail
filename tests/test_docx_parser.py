from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement

from spectrail.evidence import build_quote_match_registry
from spectrail.evidence.enricher import SourceEvidenceEnricher
from spectrail.evidence.index_builder import ensure_evidence_index
from spectrail.evidence.source_identity import canonicalize_source_cell_ids
from spectrail.extractors.reqir_extractor import ReqIRExtractor
from spectrail.llm.base import ModelRequest
from spectrail.llm.prompt_builder import build_reqir_prompt
from spectrail.parsers.docx_parser import DocxParser, DocxParserV2
from spectrail.parsers.registry import parse_document
from spectrail.validators.source_locator_validator import SourceLocatorValidator
from spectrail.validators.source_quote_validator import SourceQuoteValidator


def test_docx_parser_extracts_blocks_in_document_order(tmp_path: Path):
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_heading("System Requirements", level=1)
    document.add_paragraph("The system shall export requirements as XLSX.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Field"
    table.cell(0, 1).text = "Description"
    table.cell(1, 0).text = "timeout"
    table.cell(1, 1).text = "The system shall respond within 3 seconds."
    document.add_paragraph("The system shall allow review actions.", style="List Bullet")
    document.save(path)

    parsed = DocxParser().parse(path)

    assert parsed.document_name == "sample.docx"
    assert parsed.source_format == "docx"
    assert parsed.parser_name == "docx_parser_v2"
    assert [block.type for block in parsed.blocks] == [
        "heading",
        "paragraph",
        "table",
        "list",
    ]
    assert [block.order for block in parsed.blocks] == [1, 2, 3, 4]
    assert [block.block_id for block in parsed.blocks] == [
        "blk_0001",
        "blk_0002",
        "blk_0003",
        "blk_0004",
    ]
    assert parsed.blocks[0].section_path == ["System Requirements"]
    assert parsed.blocks[1].section_path == ["System Requirements"]
    assert parsed.blocks[2].section_path == ["System Requirements"]
    assert parsed.blocks[3].section_path == ["System Requirements"]
    assert parsed.blocks[3].section_path == ["System Requirements"]
    assert parsed.blocks[0].metadata["source_format"] == "docx"
    assert parsed.blocks[0].metadata["parser"] == "docx_parser_v2"
    assert parsed.blocks[0].metadata["style"] == "Heading 1"
    assert parsed.blocks[2].text == (
        "Field | Description\n"
        "timeout | The system shall respond within 3 seconds."
    )
    assert "##" not in parsed.text
    assert "# System Requirements" in parsed.text
    assert parsed.parser_identity is not None
    assert parsed.parser_identity.parser_name == "docx_parser_v2"
    assert parsed.parser_identity.parser_version == "2"
    assert parsed.evidence_index is not None
    assert parsed.evidence_index.tables[0].block_ids == ["blk_0003"]
    assert parsed.evidence_index.tables[0].row_count == 2
    assert parsed.evidence_index.tables[0].column_count == 2
    assert parsed.evidence_index.blocks[2].table_row_start == 1
    assert parsed.evidence_index.blocks[2].table_row_end == 2
    assert ensure_evidence_index(path, parsed) == parsed.evidence_index


def test_parse_document_routes_docx(tmp_path: Path):
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("The system shall import DOCX files.")
    document.save(path)

    parsed = parse_document(path)

    assert parsed.source_format == "docx"
    assert parsed.parser_name == "docx_parser_v2"
    assert parsed.blocks[0].text == "The system shall import DOCX files."
    assert DocxParser is DocxParserV2


def test_docx_parser_builds_merged_cell_grid_and_occurrences(tmp_path: Path):
    path = tmp_path / "merged.docx"
    document = Document()
    table = document.add_table(rows=3, cols=3)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "Header"
    table.cell(0, 2).text = "Status"
    _mark_repeating_header(table.rows[0])
    table.cell(1, 0).merge(table.cell(2, 0)).text = "Merged"
    table.cell(1, 1).text = "A"
    table.cell(1, 2).text = "B"
    table.cell(2, 1).text = ""
    table.cell(2, 2).text = "C"
    document.save(path)

    parsed = DocxParser().parse(path)
    index = parsed.evidence_index
    assert index is not None
    assert [block.text for block in parsed.blocks] == [
        "Header | Status\nMerged | A | B\nMerged |  | C"
    ]
    assert index.tables[0].row_count == 3
    assert index.tables[0].column_count == 3
    assert index.tables[0].topology_status == "complete"
    assert index.tables[0].warnings == []

    cells = {cell.cell_id: cell for cell in index.cells}
    header = cells["cell_00000001_r0001_c0001"]
    merged = cells["cell_00000001_r0002_c0001"]
    empty = cells["cell_00000001_r0003_c0002"]
    assert header.column_span == 2
    assert header.row_span == 1
    assert header.is_header is True
    assert cells["cell_00000001_r0001_c0003"].is_header is True
    assert merged.row_span == 2
    assert merged.column_span == 1
    assert empty.text == ""

    merged_occurrences = [
        item for item in index.cell_occurrences if item.cell_id == merged.cell_id
    ]
    assert [item.physical_row_index for item in merged_occurrences] == [2, 3]
    assert [item.occurrence_role for item in merged_occurrences] == [
        "original",
        "row_span_projection",
    ]
    for occurrence in index.cell_occurrences:
        block = next(
            item for item in parsed.blocks if item.block_id == occurrence.block_id
        )
        cell = cells[occurrence.cell_id]
        assert block.text[
            occurrence.canonical_start : occurrence.canonical_end
        ] == cell.text
    reparsed = DocxParser().parse(path)
    assert reparsed.evidence_index == index


def test_docx_parser_supports_legacy_horizontal_merge_markup(tmp_path: Path):
    path = tmp_path / "legacy-hmerge.docx"
    document = Document()
    table = document.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(0, 2).text = "C"
    _set_horizontal_merge(table.cell(0, 0), "restart")
    _set_horizontal_merge(table.cell(0, 1), "continue")
    document.save(path)

    parsed = DocxParser().parse(path)
    index = parsed.evidence_index
    assert index is not None
    assert parsed.blocks[0].text == "A B | C"
    assert [
        (cell.column_index, cell.column_span, cell.text)
        for cell in index.cells
    ] == [(1, 2, "A B"), (3, 1, "C")]


def test_docx_parser_groups_large_tables_and_projects_repeated_header(tmp_path: Path):
    path = tmp_path / "large.docx"
    document = Document()
    table = document.add_table(rows=22, cols=2)
    table.cell(0, 0).text = "Key"
    table.cell(0, 1).text = "Requirement"
    _mark_repeating_header(table.rows[0])
    for row_index in range(1, 22):
        table.cell(row_index, 0).text = f"R{row_index}"
        table.cell(row_index, 1).text = f"Requirement {row_index}"
    document.save(path)

    parsed = DocxParser().parse(path)
    index = parsed.evidence_index
    assert index is not None
    assert len(parsed.blocks) == 2
    assert index.tables[0].block_ids == ["blk_0001", "blk_0002"]
    assert (
        index.blocks[0].table_row_start,
        index.blocks[0].table_row_end,
    ) == (1, 20)
    assert (
        index.blocks[1].table_row_start,
        index.blocks[1].table_row_end,
    ) == (21, 22)
    assert parsed.blocks[1].text.startswith("Key | Requirement\nR20 | Requirement 20")

    header_id = "cell_00000001_r0001_c0001"
    header_occurrences = [
        item for item in index.cell_occurrences if item.cell_id == header_id
    ]
    assert [(item.block_id, item.occurrence_role) for item in header_occurrences] == [
        ("blk_0001", "original"),
        ("blk_0002", "repeated_header"),
    ]
    assert header_id in index.blocks[1].cell_ids
    assert index.tables[0].warnings == ["DOCX_REPEATED_HEADER_PROJECTED"]


def test_docx_table_serializer_escapes_content_without_collapsing_breaks(tmp_path: Path):
    path = tmp_path / "escaping.docx"
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "path | pattern\\root\nline 2"
    table.cell(0, 0).add_paragraph("paragraph 2")
    table.cell(0, 1).text = "  keep   spaces  "
    document.save(path)

    parsed = DocxParser().parse(path)
    assert parsed.blocks[0].text == (
        r"path \| pattern\\root\nline 2\nparagraph 2 |   keep   spaces  "
    )
    index = parsed.evidence_index
    assert index is not None
    assert index.cells[0].text == r"path \| pattern\\root\nline 2\nparagraph 2"
    for occurrence in index.cell_occurrences:
        cell = next(item for item in index.cells if item.cell_id == occurrence.cell_id)
        assert parsed.blocks[0].text[
            occurrence.canonical_start : occurrence.canonical_end
        ] == cell.text


def test_docx_grid_slots_without_xml_cells_downgrade_to_text_only(tmp_path: Path):
    path = tmp_path / "grid-before.docx"
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    _set_row_grid_before(table.rows[0], 1)
    document.save(path)

    parsed = DocxParser().parse(path)
    index = parsed.evidence_index
    assert index is not None
    assert parsed.blocks[0].text == "A | B"
    assert index.tables == []
    assert index.cells == []
    assert index.cell_occurrences == []
    assert index.blocks[0].table_id is None
    assert index.blocks[0].available_capabilities == ["text_range"]
    assert parsed.warnings == ["DOCX_TABLE_TOPOLOGY_UNAVAILABLE: table 1"]


def test_docx_invalid_merge_falls_back_with_warning(tmp_path: Path):
    path = tmp_path / "invalid-merge.docx"
    document = Document()
    table = document.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "A"
    table.cell(1, 0).text = "B"
    _set_vertical_merge(table.cell(0, 0), "continue")
    document.save(path)

    parsed = DocxParser().parse(path)
    index = parsed.evidence_index
    assert index is not None
    assert [cell.text for cell in index.cells] == ["A", "B"]
    assert "DOCX_MERGED_CELL_BEST_EFFORT" in parsed.warnings
    assert index.tables[0].warnings == ["DOCX_MERGED_CELL_BEST_EFFORT"]


def test_docx_table_runs_structured_evidence_validation_end_to_end(tmp_path: Path):
    path = tmp_path / "structured.docx"
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Field"
    table.cell(0, 1).text = "Requirement"
    table.cell(1, 0).text = "timeout"
    table.cell(1, 1).text = "The system shall respond within 3 seconds."
    document.save(path)

    parsed = DocxParser().parse(path)
    index = parsed.evidence_index
    assert index is not None
    prompt = build_reqir_prompt(
        ModelRequest(
            document_text=parsed.text,
            document_name=parsed.document_name,
            source_format=parsed.source_format,
            parser_name=parsed.parser_name,
            model_mode="mock",
            blocks=parsed.blocks,
            evidence_index=index,
            metadata={"evidence_policy": "structured_required"},
        )
    )
    source_cell_id = "cell_00000001_r0002_c0002"
    assert source_cell_id in prompt
    assert "row 2:" in prompt

    requirement = ReqIRExtractor().extract(
        {
            "items": [
                {
                    "statement": "The system shall respond within 3 seconds.",
                    "source_block_id": parsed.blocks[0].block_id,
                    "source_quote": "The system shall respond within 3 seconds.",
                    "source_cell_ids": [source_cell_id],
                    "source_table_row_index": 2,
                }
            ]
        },
        parsed.blocks,
        document_name=parsed.document_name,
    )[0]
    canonicalize_source_cell_ids([requirement], index)
    registry = build_quote_match_registry(
        [requirement],
        parsed.blocks,
        evidence_fingerprint=index.evidence_fingerprint,
        evidence_index=index,
    )
    SourceEvidenceEnricher().enrich(
        [requirement], index, registry, parsed.blocks
    )
    quote_validated, quote_report = SourceQuoteValidator().validate(
        [requirement], parsed.blocks, registry
    )
    locator_validated, locator_report, failures = SourceLocatorValidator().validate(
        quote_validated,
        index,
        registry,
        policy="structured_required",
        document_blocks=parsed.blocks,
    )

    assert quote_report.valid is True
    assert locator_report.valid is True
    assert failures == []
    assert locator_validated == [requirement]
    assert requirement.sources[0].locator_status == "PASS_STRUCTURED"
    assert {
        result.capability: result.status
        for result in requirement.sources[0].capability_results
    } == {"text_range": "PASS", "table_cell": "PASS"}


def _mark_repeating_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    properties.append(OxmlElement("w:tblHeader"))


def _set_horizontal_merge(cell, value: str) -> None:
    merge = OxmlElement("w:hMerge")
    merge.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val",
        value,
    )
    cell._tc.get_or_add_tcPr().append(merge)


def _set_vertical_merge(cell, value: str) -> None:
    merge = OxmlElement("w:vMerge")
    merge.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val",
        value,
    )
    cell._tc.get_or_add_tcPr().append(merge)


def _set_row_grid_before(row, value: int) -> None:
    grid_before = OxmlElement("w:gridBefore")
    grid_before.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val",
        str(value),
    )
    row._tr.get_or_add_trPr().append(grid_before)
