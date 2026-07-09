from __future__ import annotations

from pathlib import Path

import fitz
from docx import Document

from spectrail.core.io import read_json
from spectrail.parsers.markdown_parser import MarkdownParser
from spectrail.pipeline import PipelineRunner


def test_pipeline_extract_docx_mock_with_markdown_aligned_blocks(tmp_path: Path):
    source_blocks = MarkdownParser().parse_file("docs/sample_srs.md")
    document_path = tmp_path / "sample_srs.docx"
    _write_docx_from_blocks(document_path, source_blocks)

    output = tmp_path / "demo_docx"
    result = PipelineRunner().extract(document_path, output)

    manifest = read_json(result.manifest_path)
    assert manifest["status"] == "completed"
    assert manifest["counts"]["blocks"] == len(source_blocks)
    assert result.validated_count >= 14
    assert (output / "parsed" / "document.md").exists()
    assert (output / "parsed" / "blocks.json").exists()
    assert (output / "exports" / "reqir.json").exists()
    assert (output / "exports" / "requirements.xlsx").exists()


def test_pipeline_extract_text_pdf_mock_with_page_sources(tmp_path: Path):
    source_blocks = MarkdownParser().parse_file("docs/sample_srs.md")
    document_path = tmp_path / "sample_srs_text.pdf"
    _write_pdf_from_blocks(document_path, source_blocks)

    output = tmp_path / "demo_pdf"
    result = PipelineRunner().extract(document_path, output)

    manifest = read_json(result.manifest_path)
    blocks = read_json(output / "parsed" / "blocks.json")
    reqir = read_json(output / "exports" / "reqir.json")

    assert manifest["status"] == "completed"
    assert manifest["counts"]["blocks"] == len(source_blocks)
    assert result.validated_count >= 14
    assert any(item["sources"][0]["page"] is not None for item in reqir["items"])

    blocks_by_id = {block["block_id"]: block for block in blocks}
    for item in reqir["items"]:
        source = item["sources"][0]
        assert source["page"] == blocks_by_id[source["block_id"]]["page"]


def _write_docx_from_blocks(path: Path, blocks: list) -> None:
    document = Document()
    for block in blocks:
        if block.type == "heading":
            document.add_heading(block.text, level=int(block.metadata.get("level", 1)))
        elif block.type == "list":
            document.add_paragraph(block.text, style="List Bullet")
        else:
            document.add_paragraph(block.text)
    document.save(path)


def _write_pdf_from_blocks(path: Path, blocks: list) -> None:
    document = fitz.open()
    for block in blocks:
        page = document.new_page(width=2000, height=400)
        page.insert_text((72, 72), block.text, fontsize=8, fontname="china-s")
    document.save(path)
    document.close()
